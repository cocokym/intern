import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_font(cell, font_name='Calibri', font_size=12, is_bold=False):
    text = cell.text  # Store the existing text
    paragraph = cell.paragraphs[0]
    paragraph.clear()  # Clear existing formatting
    
    # Add the text back with new formatting
    run = paragraph.add_run(text)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = is_bold

def create_gene_table(doc, row_data, is_trio):
    # Create table with 4 rows and 7 columns for the main format
    table = doc.add_table(rows=4, cols=7)
    table.style = 'Table Grid'
    table.allow_autofit = False  # Disable autofit
    
    # Set column widths using cell widths
    widths = [2.11, 5.64, 1.75, 2.25, 3.0, 2.0, 2.25]  # in cm
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx])
    
    # Set overall table preferred width
    table.width = Cm(sum(widths))
    
    # Set row heights
    table.rows[0].height = Cm(1.25)  # First header row
    table.rows[1].height = Cm(1.25)  # First data row
    table.rows[2].height = Cm(0.7)   # Second header row
    table.rows[3].height = Cm(0.7)   # Second data row
    
    # First row headers
    headers_row1 = ['Gene name/OMIM', 'Transcript/Variant in HGVS Nomenclature', 'Exon Location', 
                   'Genotype/Zygosity', 'Inheritance', 'Parent origin', 'Classification']
    
    # Second row headers (shifted right by one cell)
    headers_row2 = ['', 'Position', 'REF/ALT', 'Assembly', 'SNP Identifier', 'Phenotype', '']
    
    # Merge cells for Phenotype (row 3, columns 5 and 6)
    cell_phenotype = table.cell(2, 5)
    cell_empty = table.cell(2, 6)
    cell_phenotype.merge(cell_empty)
    
    # Fill first row headers with formatting
    for i, header in enumerate(headers_row1):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        
    # Fill first row data with formatting
    row = table.rows[1]
    data = [
        f"{row_data['Gene Names']}*{row_data['OMIM ID']}",
        f"{row_data['HGVS c. (Clinically Relevant)']}\n{row_data['HGVS p. (Clinically Relevant)']}",
        str(row_data['Exon Number (Clinically Relevant)']),
        str(row_data['Zygosity']),
        str(row_data['Inheritance']),
        str(row_data['Inherited From']) if is_trio else "",  # Parent origin only for trio
        str(row_data['Second review and comment on reportable variant '])  # Note the space at the end
    ]
    
    for i, value in enumerate(data):
        cell = row.cells[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        # Make Gene name/OMIM content bold
        run.font.bold = (i == 0)  # Bold only for first column
    
    # Fill second row headers with formatting
    for i, header in enumerate(headers_row2[:-1]):  # Exclude the last empty header
        cell = table.cell(2, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        
    # Merge cells for Phenotype data (row 4, columns 5 and 6)
    cell_phenotype_data = table.cell(3, 5)
    cell_empty_data = table.cell(3, 6)
    cell_phenotype_data.merge(cell_empty_data)
    
    # Fill second row data with formatting
    row = table.rows[3]
    position_ref_alt = row_data['Chr:Pos'].split(':')
    data = [
        "",  # First cell empty
        position_ref_alt[0],  # Position
        position_ref_alt[1],  # REF/ALT
        "GRCh38",
        str(row_data['RSID']),
        str(row_data['Title']),  # Phenotype in merged cell
    ]
    
    for i, value in enumerate(data):
        cell = row.cells[i] if i < 5 else cell_phenotype_data
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = False
    
    # Add spacing after table
    doc.add_paragraph()

    # Ensure the table is not split across pages
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.page_break_before = True

def main():
    # Read the Excel files with full paths, skipping the first row and using the second row as the header
    singleton_excel_path = "/Users/admin/Desktop/bioinformatics/biof3004 intern/code/table/IMM_GRCh38_SuperPanel_554G_MANE_template_V3.1_IM662_SuperPanel_Coco_singleton.xlsx"
    trio_excel_path = "/Users/admin/Desktop/bioinformatics/biof3004 intern/code/table/IMM_GRCH38_Trio_Template_V3_IM673_674_675_CompoundHomozygous_trioformat_coco.xlsx"
    
    print(f"Reading Singleton Excel file from: {singleton_excel_path}")
    singleton_df = pd.read_excel(singleton_excel_path, header=1)
    print(f"Found {len(singleton_df)} rows in Singleton Excel file")
    print("Singleton DataFrame columns:", singleton_df.columns)
    
    print(f"Reading Trio Excel file from: {trio_excel_path}")
    trio_df = pd.read_excel(trio_excel_path, header=1)
    print(f"Found {len(trio_df)} rows in Trio Excel file")
    print("Trio DataFrame columns:", trio_df.columns)
    
    # Create new document
    doc = Document()
    
    # Create a table for each row in the Singleton Excel file
    for index, row in singleton_df.iterrows():
        print(f"Creating table for Singleton row {index + 1}")
        create_gene_table(doc, row, is_trio=False)
    
    # Create a table for each row in the Trio Excel file
    for index, row in trio_df.iterrows():
        print(f"Creating table for Trio row {index + 1}")
        create_gene_table(doc, row, is_trio=True)
    
    # Save the document
    output_path = 'table_results.docx'
    doc.save(output_path)
    print(f"Tables saved to: {output_path}")

if __name__ == "__main__":
    main()
