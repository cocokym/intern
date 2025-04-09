from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd
from datetime import datetime
import re
from docx import Document
import os
from docx.shared import Pt
import sqlite3
from db_utils import DatabaseManager
from werkzeug.utils import secure_filename

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'  # Ensure this matches the directory where files are saved
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'xlsx', 'xls'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Create uploads directory if it doesn't exist
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Global variable to store the DataFrame
df = None

# Update the database configuration
db = DatabaseManager()

# Add error handling for database connection
def load_patient_data():
    """Load patient data from database"""
    try:
        return db.get_all_patients()
    except Exception as e:
        print(f"Database connection error: {str(e)}")
        return None

def get_all_patients(self):
    try:
        with self.get_connection() as conn:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT * FROM patients")
            results = cursor.fetchall()
            print("Database results:", results)  # Debug print
            return pd.DataFrame(results) if results else pd.DataFrame()
    except Exception as e:
        print(f"Error in get_all_patients: {str(e)}")  # Debug print
        return pd.DataFrame()


def load_excel_data():
    global df
    try:
        # Read the Excel file
        df = pd.read_excel('IM patient list_20250303_template.xlsx', header=1)
        
        # Print original columns for debugging
        print("Original columns before mapping:", df.columns.tolist())
        
        # First, handle unnamed columns if they exist
        unnamed_mapping = {}
        for col in df.columns:
            if 'Unnamed' in str(col):
                col_index = int(col.split(':')[1])
                if col_index == 0:
                    unnamed_mapping[col] = 'IM Lab. no.'
                elif col_index == 1:
                    unnamed_mapping[col] = 'Lab. no.'
                elif col_index == 2:
                    unnamed_mapping[col] = 'Patient name'
                # ... add more mappings as needed
        
        if unnamed_mapping:
            df = df.rename(columns=unnamed_mapping)
        
        # Now map to our standardized names
        column_mapping = {
            'Singe gene Reported date': 'report_date',
            'Lab. no.': 'lab_number',
            'IM Lab. no.': 'im_lab_number',
            'Patient name': 'name',
            'HKID': 'hkid',
            'DOB': 'dob',
            'Ethnicity': 'ethnicity',
            'Sample collection date': 'specimen_collected',
            'Sample receive date': 'specimen_arrived',
            'Sex/Age': 'Sex/Age',
            'Case': 'case_history',  # Map 'Case' to 'case_history'
            'Type of test': 'type_of_test',
            'Type of findings': 'type_of_findings'
        }
        
        # Rename columns
        df = df.rename(columns=column_mapping)
        
        # Process Sex/Age column
        if 'Sex/Age' in df.columns:
            # Split Sex/Age and create new columns
            df[['sex', 'age']] = df['Sex/Age'].str.split('/', expand=True)
            # Clean the new columns
            df['sex'] = df['sex'].astype(str).str.strip()
            df['age'] = df['age'].astype(str).str.strip()
        
        # Convert relevant columns to string and clean them
        string_columns = ['im_lab_number', 'lab_number', 'name', 'hkid', 'ethnicity', 'clinical_history', 'type_of_test', 'type_of_findings']
        for col in string_columns:
            if col in df.columns:
                df[col] = df[col].astype(str).str.strip()
                df[col] = df[col].replace('nan', '')
        
        # Convert date columns to datetime
        date_columns = ['report_date', 'dob', 'specimen_collected', 'specimen_arrived']
        for col in date_columns:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors='coerce')
        
        # Print debug information
        print("\nFinal columns:", df.columns.tolist())
        print("\nFirst row:")
        print(df.iloc[0])
        
        return True
    except Exception as e:
        print(f"Error in load_excel_data: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return False

def validate_lab_number(lab_number):
    im_pattern = r'^IM\d{3}$'
    num_pattern = r'^2\d{10}$'
    return bool(re.match(im_pattern, lab_number) or re.match(num_pattern))

def create_word_document(patient_data):
    try:
        output_doc = Document()
        
        style = output_doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)
        
        # Add report date with label and proper formatting
        p = output_doc.add_paragraph()
        p.add_run("REPORT DATE: ").bold = True
        # Handle date formatting safely
        try:
            if patient_data['report_date']:
                date_obj = datetime.strptime(patient_data['report_date'], '%Y-%m-%d')
                formatted_date = date_obj.strftime('%d/%m/%Y')
            else:
                formatted_date = ''
        except:
            formatted_date = patient_data['report_date']
        p.add_run(formatted_date).bold = True
        
        # Add patient information with formatting
        info = [
            ('Lab. #', f"{patient_data['im_lab_number']}/{patient_data['lab_number']}"),
            ('Name', patient_data['name']),
            ('HKID', patient_data['hkid']),
            ('Date of Birth', patient_data['dob']),
            ('Sex', patient_data['sex']),
            ('Age', patient_data['age']),
            ('Ethnicity', patient_data['ethnicity']),
            ('Specimen Collected', patient_data['specimen_collected']),
            ('Specimen Arrived', patient_data['specimen_arrived'])
        ]
        
        for label, value in info:
            p = output_doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            # Format dates if the value is a date
            if label in ['Date of Birth', 'Specimen Collected', 'Specimen Arrived'] and value:
                try:
                    date_obj = datetime.strptime(value, '%Y-%m-%d')
                    value = date_obj.strftime('%d/%m/%Y')
                except ValueError:
                    pass
            p.add_run(str(value))
        
        # Add line separator at the end
        p = output_doc.add_paragraph()
        p.add_run("-" * 117)
        
        # Add summary of results
        sections = [
            ('SPECIMEN', 'EDTA blood'),
            ('CLINICAL HISTORY', patient_data['clinical_history']),  # Use clinical_history key
            ('TYPE OF TESTING REQUESTED', patient_data.get('type_of_test', '')),
            ('TEST DESCRIPTION', get_test_description(patient_data.get('test_type', ''))),
            ('SUMMARY OF RESULT(S)', get_summary_result(patient_data.get('type_of_findings', ''), patient_data.get('lab_number')))
        ]

        for label, value in sections:
            # Add title paragraph
            p = output_doc.add_paragraph()
            p.add_run(f"{label}:").bold = True

            # Add value on next line
            p = output_doc.add_paragraph()
            p.add_run(str(value))
        
        # Create filename using lab number and timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"patient_info_{patient_data['lab_number']}_{timestamp}.docx"
        
        # Save the document
        output_doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        return None

def get_test_description(test_type):
    base_desc = "In-house Immunological Disorders SuperPanel gene panel from WES was tested by next generation sequencing, and 516 genes were included in the panel test."
    if test_type.lower() == 'trio':
        return f"{base_desc} Trio analysis has been performed."
    return base_desc

def get_summary_result(finding_type, lab_number):
    print(f"Debug: Generating summary for lab number {lab_number}")  # Debug print
    file_path = db.get_uploaded_file_path(lab_number)
    if file_path:
        if not os.path.exists(file_path):
            print(f"Debug: File does not exist at path: {file_path}")  # Debug print
            return "Uploaded file not found on the server."

        try:
            # Read the uploaded file
            variant_df = pd.read_excel(file_path)

            # Filter rows where "Reportable Variant" is not empty
            reportable_variants = variant_df[variant_df['Reportable Variant'].notnull()]

            # Generate sentences for each reportable variant
            sentences = []
            for _, row in reportable_variants.iterrows():
                comment = row['Second review and comment on reportable variant ']
                gene_name = row['Gene Names']
                sentences.append(f"One likely {comment} variant was detected in the {gene_name} gene.")

            # Join all sentences into a single summary
            return " ".join(sentences)

        except Exception as e:
            print(f"Error processing file for lab number {lab_number}: {str(e)}")
            return "Error processing the uploaded file for this patient."
    else:
        print(f"Debug: No file path found for lab number {lab_number}")  # Debug print
        return "No uploaded file found for this patient."

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/search', methods=['POST'])
def search():
    global df
    if df is None:
        if not load_excel_data():
            return jsonify({
                'success': False,
                'message': 'Error loading patient database.'
            })

    lab_number = request.form.get('lab_number')
    test_type = request.form.get('test_type')
    
    if not validate_lab_number(lab_number):
        return jsonify({
            'success': False,
            'message': 'Invalid lab number format. Please use IMxxx or 2xxxxxxxxxx format.'
        })

    result = df[(df['lab_number'] == lab_number) | (df['im_lab_number'] == lab_number)]

    if not result.empty:
        patient = result.iloc[0]
        
        # Handle dates safely
        def format_date(date_value):
            if pd.isna(date_value):
                return ''
            try:
                if isinstance(date_value, str):
                    # Try to parse the string date
                    try:
                        date_obj = pd.to_datetime(date_value).date()
                        return date_obj.strftime('%Y-%m-%d')
                    except:
                        return date_value
                return date_value.date().strftime('%Y-%m-%d')
            except:
                return str(date_value)

        # Prepare patient data with safe date handling
        patient_data = {
            'report_date': format_date(patient['report_date']),
            'im_lab_number': str(patient['im_lab_number']),
            'lab_number': str(patient['lab_number']),
            'name': str(patient['name']),
            'hkid': str(patient['hkid']),
            'dob': format_date(patient['dob']),
            'sex': str(patient['sex']),
            'age': str(patient['age']),
            'ethnicity': str(patient['ethnicity']),
            'specimen_collected': format_date(patient['specimen_collected']),
            'specimen_arrived': format_date(patient['specimen_arrived']),
            'clinical_history': str(patient['case_history']),  # Use mapped column name
            'type_of_test': str(patient['type_of_test']),
            'test_type': request.form.get('test_type'),
            'type_of_findings': str(patient['type_of_findings'])
        }
        
        # Create Word document
        doc_filename = create_word_document(patient_data)
        
        if doc_filename:
            return jsonify({
                'success': True,
                'data': patient_data,
                'document': doc_filename
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Error creating Word document.'
            })
    else:
        return jsonify({
            'success': False,
            'message': 'No patient found with this lab number.'
        })

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return str(e)

# Update add_new_patient function
def add_new_patient(patient_data):
    """Add new patient using database manager"""
    return db.add_patient(patient_data)

# Add a new route for adding patients
@app.route('/add_patient', methods=['POST'])
def add_patient():
    try:
        patient_data = request.get_json()
        
        # Validate required fields
        required_fields = ['lab_number', 'im_lab_number', 'name']
        for field in required_fields:
            if not patient_data.get(field):
                return jsonify({
                    'success': False,
                    'message': f'Missing required field: {field}'
                })
        
        # Add timestamp
        patient_data['created_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Save to database
        success = db.add_patient(patient_data)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Patient added successfully'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Failed to add patient'
            })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/get_patients', methods=['GET'])
def get_patients():
    try:
        # Fetch all patients from the database
        patients = db.get_all_patients()
        if patients.empty:
            return jsonify({'success': True, 'data': [], 'filters': {}})

        # Extract unique values for dropdown filters
        filters = {
            'lab_numbers': patients['lab_number'].dropna().unique().tolist(),
            'im_lab_numbers': patients['im_lab_number'].dropna().unique().tolist(),
            'names': patients['name'].dropna().unique().tolist(),
            'test_types': patients['type_of_test'].dropna().unique().tolist(),
            'findings': patients['type_of_findings'].dropna().unique().tolist(),
        }

        return jsonify({
            'success': True,
            'data': patients.to_dict('records'),
            'filters': filters
        })
    except Exception as e:
        print(f"Error in /get_patients: {str(e)}")  # Debug print
        return jsonify({'success': False, 'message': str(e)})

@app.route('/update_findings', methods=['POST'])
def update_findings():
    try:
        data = request.get_json()
        lab_number = data.get('lab_number')
        findings = data.get('type_of_findings')

        if not lab_number or not findings:
            return jsonify({'success': False, 'message': 'Lab number and findings are required'})

        success = db.update_findings(lab_number, findings)

        if success:
            return jsonify({'success': True, 'message': 'Findings updated successfully'})
        else:
            return jsonify({'success': False, 'message': 'Failed to update findings'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/view_patient/<lab_number>', methods=['GET'])
def view_patient(lab_number):
    try:
        patient = db.get_patient(lab_number)
        if patient is not None:
            return jsonify({'success': True, 'data': patient})
        return jsonify({'success': False, 'message': 'Patient not found'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/delete_patient/<lab_number>', methods=['DELETE'])
def delete_patient(lab_number):
    try:
        success = db.delete_patient(lab_number)

        if success:
            return jsonify({'success': True, 'message': 'Patient deleted successfully'})
        else:
            return jsonify({'success': False, 'message': 'Failed to delete patient'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

def process_variant_file(file_path, lab_number):
    try:
        # Read the Excel file
        variant_df = pd.read_excel(file_path)

        # Find rows where 'Reportable Variant' is 'C'
        c_variants = variant_df[variant_df['Reportable Variant'] == 'C']

        if not c_variants.empty:
            # Get the first C variant's information
            variant = c_variants.iloc[0]
            variant_comment = variant['Second review and comment on reportable variant ']
            gene_name = variant['Gene Names']

            # Create the formatted summary
            summary = f"One likely pathogenic variant was detected in the {gene_name} gene."

            # Update the database
            success = db.update_findings_and_summary(lab_number, 'C', summary)

            if success:
                return True, 'Summary updated successfully'
            else:
                return False, 'Failed to update database'

        # Handle cases where no "C" variant is found
        return False, 'No reportable variant (C) found in file'

    except Exception as e:
        return False, f'Error processing file: {str(e)}'

@app.route('/upload_variant_file', methods=['POST'])
def upload_variant_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file uploaded'})
        
        file = request.files['file']
        lab_number = request.form.get('lab_number')
        
        # Validate inputs
        if not file or file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'})
            
        if not lab_number:
            return jsonify({'success': False, 'message': 'Lab number is required'})
            
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': 'Invalid file type. Only .xlsx and .xls files are allowed'})
        
        # Create uploads directory if it doesn't exist
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            variant_df = pd.read_excel(file_path, header=1)
            
            # Check required columns
            required_columns = ['Reportable Variant', 'Second review and comment on reportable variant ', 'Gene Names']
            missing_columns = [col for col in required_columns if col not in variant_df.columns]
            
            if missing_columns:
                return jsonify({
                    'success': False,
                    'message': f'Missing required columns: {", ".join(missing_columns)}'
                })
            
            # Get first row with any reportable variant
            variants = variant_df[variant_df['Reportable Variant'].notna()]
            
            if not variants.empty:
                variant = variants.iloc[0]
                variant_type = variant['Reportable Variant']
                
                # Set summary based on variant type
                if variant_type == 'C':
                    variant_comment = variant['Second review and comment on reportable variant ']
                    gene_name = variant['Gene Names']
                    summary = f"One {variant_comment} variant was detected in the {gene_name} gene"
                elif variant_type == 'A':
                    summary = "No disease-causing variant detected to fully account for the patient's phenotype. However, details on some additional findings have been included for reference."
                elif variant_type in ['I', 'N']:
                    summary = "No disease-causing variant detected to fully account for the patient's phenotype."
                else:
                    summary = ""
                
                # Update database
                success = db.update_findings_and_summary(lab_number, variant_type, summary)
                
                if success:
                    # Get updated patient data
                    patient_data = db.get_patient(lab_number)
                    if patient_data is not None:
                        return jsonify({
                            'success': True,
                            'message': 'Variant information updated successfully',
                            'summary': summary,
                            'data': patient_data
                        })
                
                return jsonify({
                    'success': True,
                    'message': 'Variant information updated successfully',
                    'summary': summary
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'No variant found in file'
                })
                
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)
                
    except pd.errors.EmptyDataError:
        return jsonify({
            'success': False,
            'message': 'The uploaded file is empty'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error processing file: {str(e)}'
        })

@app.route('/upload_file', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file part'})

        file = request.files['file']
        file_type = request.form.get('file_type')

        if file.filename == '':
            return jsonify({'success': False, 'message': 'No selected file'})

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)

            # Save the file to the uploads folder
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            print(f"Debug: Saving file to {file_path}")  # Debug print
            file.save(file_path)

            # Extract lab number from the full file path
            lab_number_match = re.search(r'IM\d+(_\d+)*', file_path)
            if not lab_number_match:
                return jsonify({'success': False, 'message': 'Lab number not found in file path'})
            lab_number = lab_number_match.group(0)  # Fetch the full lab number (e.g., IM_651 or IM_651_652_653)

            print(f"Debug: Lab number extracted from file path: {lab_number}")  # Debug print

            # Check if a file for this lab number already exists in the uploaded_files table
            existing_file = db.get_uploaded_file_path(lab_number)
            if existing_file:
                return jsonify({
                    'success': False,
                    'message': f'A file for lab number {lab_number} already exists: {existing_file}'
                })

            # Save file information to the database
            success = db.save_uploaded_file(file_type, filename, lab_number)
            if success:
                print(f"Debug: File information saved to database for lab number {lab_number}")  # Debug print

                # Process the file and store its data in the respective table
                process_success = process_file_data(file_path, file_type, lab_number)
                if process_success:
                    return jsonify({'success': True, 'message': 'File uploaded and processed successfully', 'lab_number': lab_number})
                else:
                    return jsonify({'success': False, 'message': 'File uploaded but failed to process data'})
            else:
                return jsonify({'success': False, 'message': 'Failed to save file information to the database'})
        else:
            return jsonify({'success': False, 'message': 'Invalid file type'})
    except Exception as e:
        print(f"Error in /upload_file: {str(e)}")  # Debug print
        return jsonify({'success': False, 'message': f'Error uploading file: {str(e)}'})

@app.route('/get_uploaded_files', methods=['GET'])
def get_uploaded_files():
    try:
        with db.get_connection() as conn:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT file_type, file_name, upload_date FROM uploaded_files")
            files = cursor.fetchall()
            print("Fetched uploaded files:", files)  # Debug print
            return jsonify({'success': True, 'files': files})
    except Exception as e:
        print(f"Error in /get_uploaded_files: {str(e)}")  # Debug print
        return jsonify({'success': False, 'message': str(e)})

def delete_patient(self, lab_number):
    """Delete patient by lab number"""
    try:
        conn = self.get_connection()
        cursor = conn.cursor()
        query = "DELETE FROM patients WHERE lab_number = %s"
        cursor.execute(query, (lab_number,))
        conn.commit()
        return cursor.rowcount > 0
    except Exception as e:
        print(f"Error deleting patient: {e}")
        return False
    finally:
        if 'conn' in locals():
            conn.close()

def save_uploaded_file(self, file_type, file_name, lab_number):
    """
    Save uploaded file information to the database.
    """
    try:
        with self.get_connection() as conn:
            cursor = conn.cursor()
            query = """
                INSERT INTO uploaded_files (file_type, file_name, lab_number)
                VALUES (%s, %s, %s)
            """
            cursor.execute(query, (file_type, file_name, lab_number))
            conn.commit()
            return True
    except Exception as e:
        print(f"Error saving uploaded file: {str(e)}")
        return False

def process_file_data(file_path, file_type, lab_number):
    """
    Process the uploaded file and store its data in the respective table.
    """
    try:
        # Read the Excel file with the correct header row
        print(f"Debug: Reading file {file_path}")  # Debug print
        df = pd.read_excel(file_path, header=1)  # Use header=1 to read the second row as column headers

        # Normalize column names by stripping leading/trailing spaces
        df.columns = df.columns.str.strip()

        # Debug print for DataFrame columns
        print(f"Debug: Columns in the file after stripping spaces: {df.columns.tolist()}")

        # Replace all missing values (NaN, empty strings, whitespace) with None
        df = df.replace({pd.NA: None, '': None, ' ': None, 'nan': None, 'NaN': None}).where(pd.notnull(df), None)

        # Debug print for DataFrame after replacing missing values
        print(f"Debug: DataFrame after replacing missing values:\n{df.head()}")

        # Define the required columns for both singleton and trio tables
        required_columns = [
            "Reportable Variant", "Chr:Pos", "IGV review ( True / False call)", 
            "Second review and comment on reportable variant",  # No trailing space in the code
            "Gene Names", "HGVS c. (Clinically Relevant)", "HGVS p. (Clinically Relevant)",
            "Exon Number (Clinically Relevant)", "Zygosity", "Inheritance", "Classification",
            "OMIM ID", "RSID", "Title"
        ]

        # Add "Inherited From" for trio files
        if file_type == "trio":
            required_columns.append("Inherited From")

        # Ensure the required columns exist in the DataFrame
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            print(f"Error: Missing required columns: {', '.join(missing_columns)}")
            return False

        # Filter the DataFrame to include only the required columns
        df = df[required_columns]

        # Prepare data for insertion
        rows = []
        for _, row in df.iterrows():
            rows.append(tuple(row[col] for col in required_columns))

        # Debug print for rows
        print(f"Debug: Prepared rows for insertion: {rows}")

        # Escape column names with backticks
        escaped_columns = [f"`{col}`" for col in required_columns]

        # Insert data into the respective table
        table_name = 'singleton' if file_type == 'singleton' else 'trio'
        print(f"Debug: Inserting data into {table_name} table")  # Debug print
        with db.get_connection() as conn:
            cursor = conn.cursor()
            query = f"""
                INSERT INTO {table_name} ({', '.join(escaped_columns)})
                VALUES ({', '.join(['%s'] * len(required_columns))})
            """
            cursor.executemany(query, rows)
            conn.commit()
            print(f"Debug: Data inserted into {table_name} table for lab number {lab_number}")
        return True
    except Exception as e:
        print(f"Error processing file data: {str(e)}")
        return False

if __name__ == '__main__':
    # Load data from database when starting
    df = load_patient_data()
    if df is None:
        print("Failed to load from database, trying Excel file...")
        if not load_excel_data():
            print("Failed to load data from both database and Excel!")
    else:
        print("Successfully loaded data from database")
    
    app.run(debug=True)
